from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = Path(__file__).with_name("AI-COS_Pharmacy_Change_Log.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(85, 85, 85)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_rtl(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    p_pr.append(bidi)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_inches):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")
    cell.width = Inches(width_inches)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_paragraph(doc, text="", size=11, bold=False, color=None, after=6, align=WD_ALIGN_PARAGRAPH.RIGHT):
    paragraph = doc.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.1
    set_rtl(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return paragraph


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_width(cell, widths[index])
        set_cell_margins(cell)
        shade_cell(cell, LIGHT_BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_rtl(paragraph)
        run = paragraph.add_run(header)
        set_run_font(run, size=9.5, bold=True, color=DARK_BLUE)

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cell = cells[index]
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_rtl(paragraph)
            run = paragraph.add_run(value)
            set_run_font(run, size=9.5)
    return table


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    normal.font.size = Pt(11)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl(header)
    header_run = header.add_run("AI-COS Pharmacy | سجل التعديلات")
    set_run_font(header_run, size=9, color=GRAY)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_rtl(footer)
    footer_run = footer.add_run("وثيقة عمل داخلية — تُحدَّث بعد كل دفعة تطوير")
    set_run_font(footer_run, size=8.5, color=GRAY)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)
    set_rtl(title)
    title_run = title.add_run("سجل تعديلات مشروع AI-COS Pharmacy")
    set_run_font(title_run, size=23, bold=True, color=BLUE)

    subtitle = add_paragraph(
        doc,
        "مرجع مستمر لتوثيق التعديلات التقنية، أسبابها، ونتائج التحقق.",
        size=11,
        color=GRAY,
        after=14,
    )

    add_table(
        doc,
        ["الحالة", "آخر تحديث", "الإصدار", "اسم المشروع"],
        [["مفتوح للتحديث المستمر", str(date.today()), "0.8", "AI-COS Pharmacy"]],
        [1.55, 1.4, 0.75, 2.8],
    )

    add_paragraph(doc, "نطاق السجل", size=16, bold=True, color=BLUE, after=6)
    add_paragraph(
        doc,
        "يُسجَّل هنا كل تعديل تم على ملفات المشروع، مع توضيح السبب، الملفات المتأثرة، وحالة الاختبار. لا تُكتب المفاتيح أو كلمات المرور أو أي بيانات حساسة داخل هذا المستند.",
        after=10,
    )

    add_paragraph(doc, "التحديث رقم 1 — تأمين الإعدادات والأسرار", size=13, bold=True, color=BLUE, after=6)
    add_table(
        doc,
        ["الحالة", "التحقق", "التعديل المنفذ", "الملفات / الأثر"],
        [
            [
                "مكتمل",
                "تم التحقق من وجود 296 ملفًا في النسخة الاحتياطية (22.8 MB).",
                "إنشاء نسخة احتياطية قبل أي تعديل.",
                "backups/ai-cos-pharmacy-pre-security-20260809-01",
            ],
            [
                "مكتمل",
                "نجح Next.js production build بعد التعديل.",
                "إزالة مفتاح JWT الثابت؛ أصبح التطبيق يطلب AUTH_SESSION_SECRET عشوائيًا بطول 32 حرفًا على الأقل.",
                "frontend/src/lib/auth.ts",
            ],
            [
                "مكتمل",
                "تمت مراجعة قالب الإعداد ولم يُضمَّن فيه أي سر حقيقي.",
                "إضافة AUTH_SESSION_SECRET إلى قالب إعداد الواجهة.",
                "frontend/.env.local.example",
            ],
            [
                "مكتمل",
                "تم إنشاء قواعد تجاهل لملفات البيئة والبيانات المحلية والكاش.",
                "إضافة سياسة منع تتبع الأسرار للـBackend وقالب إعداد آمن.",
                "backend/.gitignore و backend/.env.example",
            ],
            [
                "بانتظار إجراء خارجي",
                "يتطلب دخول مالك المشروع إلى حسابات الخدمات السحابية.",
                "تدوير كلمة مرور قاعدة Supabase ومفتاح Gemini الحاليين، ثم تحديث backend/.env محليًا فقط.",
                "Supabase و Google AI Studio",
            ],
            [
                "مكتمل",
                "نجح اختبار اتصال قاعدة البيانات (SELECT 1)، واستجابة Supabase Auth، والتحقق من مفتاح Gemini دون كشف أي قيمة حساسة.",
                "التحقق من تدوير إعدادات Supabase وGemini وإعداد AUTH_SESSION_SECRET.",
                "backend/.env و frontend/.env.local",
            ],
        ],
        [0.8, 1.55, 2.65, 1.5],
    )

    add_paragraph(doc, "مراجعة ما بعد تحديثات الباك إند — 2026-08-11", size=13, bold=True, color=BLUE, after=6)
    add_table(
        doc,
        ["الحالة", "التحقق", "النتيجة", "النطاق / الأثر"],
        [
            [
                "مراجعة فقط",
                "تمت مراجعة مسارات الباك إند، الهجرات، إعدادات الصلاحيات، والواجهة دون تعديل كود المشروع.",
                "تأكيد تحسن مهم: JWT محلي عبر JWKS، صلاحيات أدوار على عدة مسارات إدارية، مفاتيح للخدمات الداخلية، وسجل أحداث قابل للتدقيق.",
                "backend/app و alembic/versions",
            ],
            [
                "مشكلة حرجة متبقية",
                "فحص قاعدة البيانات: اتصال التطبيق يستخدم مالك الجداول، وFORCE RLS غير مفعل على عدة جداول، وجدول tenant_settings بلا RLS.",
                "سياسات RLS لا تعزل كل البيانات فعلياً عند استخدام اتصال التطبيق الحالي؛ يلزم إصلاح قبل اعتبار تعدد الصيدليات آمناً.",
                "RLS / PostgreSQL multi-tenancy",
            ],
            [
                "مشكلة حرجة متبقية",
                "مفتاح N8N_SERVICE_KEY غير مضبوط في ملف التشغيل الحالي، فيعود التطبيق للقيمة الافتراضية الموجودة في الكود.",
                "المسارات الداخلية المحمية بمفتاح الخدمة ليست جاهزة للنشر الآمن حتى ضبط مفتاح عشوائي قوي خارج الكود.",
                "backend/.env و app/core/config.py",
            ],
            [
                "اختبارات جزئية",
                "نجح تجميع Python، وAlembic عند الرأس phase8_chain_per_session_fix، ونجح Next.js production build.",
                "فشل pytest قبل تنفيذ الاختبارات لأن قاعدة pharmacy_test_db على المنفذ 55432 غير مشغلة. وفشل ESLint للواجهة: 65 خطأ و24 تحذيراً.",
                "backend/tests و frontend",
            ],
            [
                "مشكلة وظيفية متبقية",
                "طلب الشراء من الواجهة يرسل items فقط، بينما API يطلب channel إلزامياً.",
                "تأكيد الطلب سيُرفض بالتحقق 422 حتى يُوحَّد عقد الواجهة والباك إند.",
                "frontend checkout / backend order schema",
            ],
        ],
        [0.95, 1.7, 2.8, 1.05],
    )

    add_paragraph(doc, "مراجعة تحقق ثانية بعد إصلاحات الفريق — 2026-08-11", size=13, bold=True, color=BLUE, after=6)
    add_table(
        doc,
        ["الحالة", "التحقق", "النتيجة", "النطاق / الأثر"],
        [
            [
                "تم التأكد من إصلاحه",
                "فحص حي لقاعدة البيانات ومسارات HTTP غير الموثقة.",
                "دور التطبيق الحالي لا يملك BYPASSRLS، وRLS مفعّل لإعدادات الصيدلية. Governance وInventory ورفع الملفات والوكلاء يرفضون الطلب غير الموثق (401). كما أن CORS يسمح بالواجهة المحلية ويرفض origin خارجي.",
                "قاعدة البيانات ومسارات API",
            ],
            [
                "خلل حرج جديد / متبقٍ",
                "عند اتصال التطبيق الحالي: current_user_tenant_id() وapp.current_tenant_id يرجعان NULL، ولا تظهر أي سجلات customers للدور التطبيقي قبل ضبط السياق.",
                "هناك حلقة منطقية: التطبيق يحتاج قراءة Customer ليضبط tenant context، لكن سياسة RLS تمنع القراءة/الإنشاء قبل ضبطه. سيؤثر على تسجيل الدخول، إنشاء Session، إنشاء Customer، وترقية الدور. سياسة customers الحالية SELECT فقط وسياسة drugs SELECT فقط رغم وجود عمليات كتابة في الـAPI.",
                "RLS + المصادقة + إدارة الدواء",
            ],
            [
                "خلل حرج متبقٍ",
                "ما زال tenant يحدد من X-Tenant-ID في إنشاء Session والتسجيل التلقائي، بينما OAuth callback لا يرسله للباك إند.",
                "يمكن اختيار tenant من العميل عند أول Session، وقد يفشل تسجيل OAuth. يلزم مصدر tenant موثوق (دعوة / نطاق / claim) لا Header يتحكم فيه العميل.",
                "Auth / multi-tenancy",
            ],
            [
                "خلل أمان متبقٍ",
                "N8N_SERVICE_KEY موجود لكنه لا يحقق معيار الطول الآمن ولا يبدو قيمة إنتاجية قوية.",
                "إزالة القيمة الافتراضية من الكود خطوة صحيحة، لكن يجب وضع مفتاح عشوائي قوي لا يقل عن 32 حرفًا والتحقق من ذلك عند بدء التطبيق.",
                "backend/.env و config",
            ],
            [
                "اختبارات فاشلة",
                "شُغلت قاعدة الاختبار فعليًا ثم pytest؛ فشلت الاختبارات الـ14 عند إنشاء جدول events.",
                "النموذج يعرّف event_seq كـ Identity مع nullable=True فينتج SQL متعارضًا (IDENTITY NULL). كذلك تشغيل compose كامل يتعطل بسبب تعارض المنفذ 6380 لـRedis. Next.js build يفشل حاليًا بسبب 6 أخطاء TypeScript، وESLint ما زال 65 خطأ و24 تحذيرًا.",
                "backend/tests و frontend",
            ],
            [
                "مخاطرة نشر",
                "دور aicos_app وسياسات RLS الخاصة بـtenant_settings موجودة في قاعدة البيانات الحية لكن لم يُعثر على migration أو سكربت مصدر لها ضمن المشروع.",
                "نشر قاعدة جديدة أو إعادة بناء البيئة لن يعيد هذه إعدادات الحماية. كما أن سياسة tenant_settings تسمح بالوصول عند NULL tenant context، وهي fail-open.",
                "Alembic / reproducibility",
            ],
        ],
        [0.95, 1.7, 2.8, 1.05],
    )

    add_paragraph(doc, "مراجعة تدقيق شاملة بعد دفعة التطوير الأخيرة — 2026-08-22", size=13, bold=True, color=BLUE, after=6)
    add_table(
        doc,
        ["الحالة", "التحقق", "النتيجة", "النطاق / الأثر"],
        [
            [
                "تم التأكد من إصلاحه",
                "تشغيل Next.js production build واختبار استجابة صفحات المتجر ولوحة الإدارة وOpenAPI.",
                "نجح بناء الواجهة، والصفحة الرئيسية و13 مساراً أساسياً تردّ 200. كما أن أعمدة الشحن والدفع للطلبات، وأعمدة البيانات السريرية للأدوية، موجودة فعلياً في قاعدة البيانات.",
                "Frontend build / قاعدة البيانات",
            ],
            [
                "تراجع أمني حرج",
                "قراءة OpenAPI واختبار HTTP دون توثيق لمسارات حساسة.",
                "60 من 71 عملية API لا تعلن حماية Authentication، منها 41 عملية كتابة. /api/v1/orders/all و/internal/governance/pending يردان 200 بدون تسجيل دخول. كما أن CORS يقبل أي Origin بصيغة http://أي-نطاق:3000.",
                "API Authorization / CORS / n8n",
            ],
            [
                "تراجع أمني حرج",
                "فحص route handlers ودالة الصلاحيات وخدمة رفع الملفات.",
                "التقييد السابق للصلاحيات غير موجود حالياً: require_role لا يتحقق من الدور، وGovernance وAgents وAnalytics ورفع الملفات والمسارات الداخلية بلا توثيق. رفع الملفات يعيد استخدام اسم العميل مباشرة دون منع traversal أو MIME/حجم أو اسم عشوائي. مفتاح N8N موجود في البيئة لكنه غير معرّف في Settings ولا يُستخدم لحماية المسارات الداخلية.",
                "RBAC / Files / Internal APIs",
            ],
            [
                "تراجع أمني حرج",
                "اتصال قراءة فقط بقاعدة التشغيل الحالية وفحص RLS والسياسات.",
                "حساب اتصال الباك إند الحالي هو postgres مع BYPASSRLS؛ لذا RLS لا يعزل البيانات من التطبيق. tenant context كان NULL. سياسة global_read_drug_interactions موجودة، لكن global_read_drug_affinities مفقودة رغم أن Alembic عند الرأس p10critfixes؛ أي أن إصلاح الهجرة منفذ جزئياً وغير قابل للإعادة تلقائياً.",
                "Supabase / RLS / Migrations",
            ],
            [
                "اختبارات فاشلة",
                "تشغيل pytest داخل PostgreSQL اختبار معزول بعد تشغيل test_db فقط وتنظيفه بعد ذلك.",
                "النتيجة: 1 ناجح، 8 فاشلة، 5 أخطاء. اختبارات Auth تستدعي signup/login/me غير الموجودة حالياً (404)، اختبارات settings تفتقد fixtures، واختبارات event chain تتطلب payload_hash غير المعرّف في ORM عند إنشاء قاعدة الاختبار. script التشغيل الكامل نفسه يتعطل عند منفذ Redis 6380 ويضبط REDIS_URL على 6389 المخالف للـcompose.",
                "backend/tests / Docker test setup",
            ],
            [
                "خلل وظيفي ونشر",
                "فحص عقود الواجهة والـAPI وnpm dependency tree.",
                "لوحة الإدارة تطلب /api/v1/auth/me لكنه غير موجود (404). ESLint يفشل بـ101 خطأ و66 تحذيراً. البناء الحالي ينجح فقط لوجود packages محلية غير مسجلة في package.json (Supabase/Prisma/Jose/bcrypt/Zustand وغيرها)، لذلك تثبيت نظيف أو CI لن يكون موثوقاً. يوجد أيضاً secret ثابت داخل src/lib/auth.ts رغم وجود AUTH_SESSION_SECRET في البيئة.",
                "Admin UI / Frontend release readiness",
            ],
            [
                "ملاحظة أكاديمية",
                "مراجعة واجهة AI Review مقابل متطلبات PRD.",
                "تنويه طبي ظاهر في chatbot، وهي نقطة صحيحة. لكن بيانات SHAP/الثقة في صفحة AI Review تُنشأ بصيغ ثابتة في الواجهة بدلاً من ناتج موثق من نموذج أو endpoint؛ لا ينبغي عرضها للجنة كدليل على Explainability فعلية حتى تُربط بنتائج النموذج وMAE/تجربة قابلة للتكرار.",
                "FR-06 / FR-12 / العرض الأكاديمي",
            ],
        ],
        [0.95, 1.7, 2.8, 1.05],
    )

    add_paragraph(doc, "مراجعة تشغيل حي للنسخة العاملة على localhost:3000 — 2026-08-22", size=13, bold=True, color=BLUE, after=6)
    add_table(
        doc,
        ["الحالة", "التحقق", "النتيجة", "النطاق / الأثر"],
        [
            [
                "تصحيح نطاق المراجعة",
                "فحص process الذي يستمع على المنفذ 3000 ثم مقارنة hashes للملفات.",
                "الموقع الحي يعمل من AI-COS-Pharmacy/frontend وليس من stitch_ai_cos_pharmacy/ai-cos-frontend. توجد نسختان متباعدتان من الواجهة والباك إند، لذلك قد يُصلح الفريق ملفاً لا تستخدمه النسخة المفتوحة في المتصفح.",
                "هيكل المشروع / سبب أخطاء متكررة",
            ],
            [
                "مانع نشر Production",
                "تشغيل npm run build من مجلد الواجهة الذي يخدم localhost:3000.",
                "فشل TypeScript في صفحة Profile: يتم تمرير drug_id إلى CartItem رغم أن النوع لا يعرّفه. Dev server يعرض الصفحة، لكن production build يفشل وبالتالي لا يمكن نشر النسخة الحالية.",
                "frontend/src/app/(storefront)/profile/page.tsx و src/lib/store/cart.ts",
            ],
            [
                "أخطاء وظائف مؤكدة",
                "مطابقة استدعاءات Agentic AI مع OpenAPI وتجربة طلبات آمنة بلا بيانات حقيقية.",
                "ثلاثة أزرار ترجع 404 لأن الواجهة تطلب analytics/kpi-summary وinventory/analyze وfinance/evaluate-discount، بينما الباك إند يوفّر أسماء مختلفة. زر Pricing يرجع 422 لأن الواجهة ترسل total_orders بدل الحقل المطلوب order_count. حقول إضافية في Marketing لا يستخدمها الباك إند.",
                "frontend/admin/agentic-ai + backend/domains/agents",
            ],
            [
                "وظيفة وهمية",
                "مراجعة صفحة Settings والـrouter المقابل.",
                "Settings في الواجهة عناصر ثابتة؛ المفاتيح لا تغير state وزر Save بلا handler. يوجد Settings router في الباك إند لكنه غير مسجل في main.py، ولذلك /api/v1/settings/ يرجع 404.",
                "Admin Settings / Backend router registration",
            ],
            [
                "اختبارات غير جاهزة",
                "جمع 35 اختباراً وتشغيلها مع PostgreSQL اختبار معزول؛ ثم فحص lastfailed.",
                "بعد التشغيل، 29 اختباراً مسجل كفاشل. تشمل اختبارات regression للأدوار وinternal endpoints، واختبارات Auth القديمة، وevent chain. كما أن compose الكامل ما زال يتعارض على Redis 6380، والسكربت يضبط REDIS_URL على 6389 رغم أن compose يعرّض 6380.",
                "backend/tests / docker-compose.test.yml",
            ],
            [
                "تراجع أمان باقٍ",
                "فحص OpenAPI، كود الحماية، واتصال قاعدة التشغيل.",
                "42 عملية كتابة بلا Authentication في OpenAPI. require_role لا يفحص الدور، orders/all وGovernance وAgents ومسارات داخلية حساسة غير محمية. كود Internal token موجود لكنه غير مستخدم ولا يمكنه العمل لأن N8N_SERVICE_KEY غير معرف في Settings. اتصال التطبيق ما زال postgres مع BYPASSRLS وtenant context فارغ.",
                "Authorization / RLS / n8n",
            ],
            [
                "تحسين مؤكد",
                "اختبار API الحالي دون token.",
                "مسارا auth/me وauth/invite-customer أصبحا موجودين الآن (الأول يرجع 401 دون token)، لكن invite-customer يتحقق من تسجيل الدخول فقط ولا يفرض دور admin.",
                "Authentication contract",
            ],
            [
                "جاهزية حزم منخفضة",
                "npm ls وnpm run lint في نسخة التشغيل.",
                "ESLint يفشل بـ102 خطأ و66 تحذيراً. Supabase SSR وZustand مستخدمان لكنهما extraneous وغير مسجلين في package.json؛ بناء بيئة نظيفة أو CI قد يفشل حتى بعد إصلاح TypeScript.",
                "Frontend dependencies / CI",
            ],
        ],
        [0.95, 1.7, 2.8, 1.05],
    )

    add_paragraph(doc, "إعادة فحص حي بعد إصلاحات إضافية — 2026-08-28", size=13, bold=True, color=BLUE, after=6)
    add_table(
        doc,
        ["الحالة", "التحقق", "النتيجة", "النطاق / الأثر"],
        [
            [
                "تحسن مؤكد",
                "تشغيل الموقع عبر المتصفح، npm run lint، وفحص فرق الملفات.",
                "الصفحة الرئيسية والكتالوج يعملان ويعرضان 114 دواءً، ولا تظهر أخطاء Console في المسارات العامة. ESLint أصبح ناجحاً، ودالة require_role تتحقق الآن من الدور فعلياً. كما أضيف تنظيف اسم الملف لمنع path traversal.",
                "Storefront / Lint / RBAC helper / Files",
            ],
            [
                "مانع نشر جديد",
                "تشغيل npm run build على نسخة الواجهة التي تخدم localhost:3000.",
                "اكتمل TypeScript لكن فشل توليد Production عند /catalog لأن StorefrontPageClient يستخدم useSearchParams خارج Suspense. الصفحة الرئيسية تغلفه بـSuspense بينما صفحة الكتالوج لا تفعل ذلك.",
                "frontend/src/app/(storefront)/catalog/page.tsx",
            ],
            [
                "خطأ حي مؤكد",
                "POST آمن إلى /api/v1/agents/analytics/summary ومراجعة توقيع الدالة.",
                "المسار يرد 500: router يرسل churn_rate وrevenue_growth بينما analyze_kpis يتطلب cross_sell_attach_rate وchurn_risk_customers وprediction_mae. كذلك ثلاثة أزرار في Agentic AI ما زالت تطلب مسارات غير موجودة، وزر Pricing يرسل total_orders بدلاً من order_count.",
                "Agents API contract / Admin Agentic AI",
            ],
            [
                "تراجع أمان حرج مستمر",
                "فحص OpenAPI وطلبات HTTP بلا token واختبار CORS من origin خارجي.",
                "42 عملية كتابة بلا Authentication. /api/v1/orders/all و/internal/governance/pending يكشفان بيانات حية دون دخول. المسارات الداخلية لا تستخدم verify_internal_token رغم وجود المفتاح في البيئة، وCORS يقبل http://any-domain:3000 مع credentials.",
                "Authorization / Internal APIs / CORS",
            ],
            [
                "إعدادات غير موصولة",
                "فحص /api/v1/settings/ ومراجعة main.py والواجهة.",
                "يوجد Settings router محمي بأدوار لكنه غير مسجل في main.py ولذلك يرجع 404؛ صفحة Settings واجهة ثابتة وزر Save بلا handler.",
                "Settings backend + Admin Settings UI",
            ],
            [
                "اختبارات غير قابلة للاعتماد",
                "نجح جمع 35 اختباراً. حاولت تشغيل المجموعة، لكن Docker daemon غير متاح وبالتالي قاعدة PostgreSQL الاختبارية على 55432 لا يمكن تجهيزها. cache السابق ما زال يحفظ 29 اختباراً فاشلاً، بينها اختبارات security regression وsettings وauth.",
                "backend/tests / Docker test environment",
            ],
            [
                "مشكلة نشر إضافية",
                "npm ls وpackage-lock في الواجهة.",
                "@supabase/ssr وzustand مستخدمان محلياً لكنهما غير مسجلين لا في package.json ولا package-lock؛ npm ci أو CI نظيف سيفشل. npm audit لإنتاج الحزم أعلن 0 ثغرات حالياً.",
                "Frontend dependencies / CI",
            ],
            [
                "تحسين جزئي فقط",
                "فحص config والمسارات الداخلية ورفع الملفات.",
                "N8N_SERVICE_KEY أصبح موجوداً في البيئة لكنه غير معرّف ضمن Settings، لذا verify_internal_token سيفشل إن تم تركيبه. تنظيف الاسم يمنع traversal فقط؛ رفع الملفات ما زال بلا دخول أو فحص نوع/حجم أو اسم فريد، وبالتالي يمكن الاستبدال أو استهلاك التخزين.",
                "n8n / File upload hardening",
            ],
        ],
        [0.95, 1.7, 2.8, 1.05],
    )

    add_paragraph(doc, "إعادة تحقق بعد دفعة إصلاحات الحماية والـ Agentic AI — 2026-08-28", size=13, bold=True, color=BLUE, after=6)
    add_table(
        doc,
        ["الحالة", "التحقق", "النتيجة", "النطاق / الأثر"],
        [
            [
                "تم إصلاحه",
                "تشغيل الموقع الحي وnpm run build ومراجعة صفحة الكتالوج.",
                "تم حل خطأ useSearchParams السابق: صفحة /catalog أصبحت تغلف StorefrontPageClient بـSuspense، والصفحة تعمل في المتصفح دون Console errors.",
                "Catalog / Storefront",
            ],
            [
                "مانع نشر جديد",
                "npm run build على نسخة الواجهة الحالية.",
                "البناء يفشل قبل TypeScript بسبب JSX غير صالح في Admin Settings: قيم className كتبت مثل w-10 h-6 وw-4 h-4 دون strings أو template literals في الأسطر 87 و90 و100 و103.",
                "frontend/src/app/admin/settings/page.tsx",
            ],
            [
                "تم إصلاحه جزئياً",
                "اختبار HTTP دون token ومراجعة مسارات FastAPI وOpenAPI.",
                "orders/all وsettings وAgent Analytics ترفض المستخدم غير الموثق (401). Governance pending يطلب X-Internal-Token، وتمت إضافة N8N_SERVICE_KEY إلى Settings. عدد عمليات الكتابة غير المعلنة في OpenAPI هبط من 42 إلى 24.",
                "Auth / Settings / Agent APIs / Internal Governance",
            ],
            [
                "خلل أمان متبقٍ",
                "مراجعة Governance internal router وOpenAPI.",
                "مسار PATCH /internal/governance/{id}/status ما زال بلا verify_internal_token، رغم حماية pending وapprove. توجد أيضاً عمليات تعديل أخرى غير محمية مثل files/upload وdrugs وwebhooks وpredictions وanalytics A/B.",
                "Internal API / Remaining authorization",
            ],
            [
                "تم إصلاحه",
                "اختبار Origin خارجي على health ومراجعة CORS.",
                "لم يعد Origin خارجي بصيغة http://any-domain:3000 يحصل على Access-Control-Allow-Origin أو credentials؛ تم استبدال regex الواسع بقائمة Origins مسموحة.",
                "CORS",
            ],
            [
                "تم إصلاحه",
                "مطابقة الواجهة مع router الجديد لوكلاء الذكاء الاصطناعي.",
                "Analytics يستخدم المسار الصحيح ويحسب بياناته من DB، وInventory وPricing تستخدم أسماء endpoint وحقول payload الصحيحة. Finance يستخدم الآن finance/clv لكن يعتمد على البحث بالاسم؛ يجب لاحقاً استبداله بمعرف العميل لتفادي تشابه الأسماء.",
                "Agentic AI contracts",
            ],
            [
                "تحسن حزم",
                "npm ls وnpm audit --omit=dev.",
                "@supabase/ssr وzustand أصبحا مسجلين في المشروع، وnpm audit للإنتاج رجع 0 ثغرات. لكن npm run lint يفشل حالياً (183 diagnostic lines) بسبب scripts مساعدة في جذر frontend وأخطاء React في صفحات Login/Profile وغيرها.",
                "Dependencies / Lint",
            ],
            [
                "خلل بيانات حرج متبقٍ",
                "اتصال قراءة فقط بقاعدة تشغيل Supabase الحالية.",
                "الباك إند ما زال يتصل بدور postgres مع rolbypassrls=true؛ لذلك RLS لا يوفر عزل tenant فعلياً من اتصال التطبيق.",
                "Supabase RLS / Multi-tenancy",
            ],
        ],
        [0.95, 1.7, 2.8, 1.05],
    )

    add_paragraph(doc, "قاعدة التحديث التالية", size=13, bold=True, color=BLUE, after=6)
    add_paragraph(
        doc,
        "بعد كل دفعة عمل، يُضاف صف جديد يوضح: التاريخ، التغيير، الملفات المتأثرة، سبب التغيير، ونتيجة الاختبار أو ما تبقى للتنفيذ.",
        after=0,
    )

    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
    print(OUTPUT_PATH)
